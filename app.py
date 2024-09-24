import streamlit as st
from streamlit_drawable_canvas import st_canvas
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.mime.text import MIMEText
import requests
import os
import socket
from datetime import datetime, date
import numpy as np
import re
from PIL import Image as PILImage
from docx.shared import Inches
import time


# Set the page configuration for the Streamlit app
st.set_page_config(
    page_title="Digital Media Release Consent Form", 
    page_icon="https://lirp.cdn-website.com/d8120025/dms3rep/multi/opt/social-image-88w.png", 
    layout="centered"
)

if 'submission_status' not in st.session_state: st.session_state.submission_status = False

def is_valid_email(email):
    # Comprehensive regex for email validation
    pattern = r'''
        ^                         # Start of string
        (?!.*[._%+-]{2})          # No consecutive special characters
        [a-zA-Z0-9._%+-]{1,64}    # Local part: allowed characters and length limit
        (?<![._%+-])              # No special characters at the end of local part
        @                         # "@" symbol
        [a-zA-Z0-9.-]+            # Domain part: allowed characters
        (?<![.-])                 # No special characters at the end of domain
        \.[a-zA-Z]{2,}$           # Top-level domain with minimum 2 characters
    '''
    
    # Match the entire email against the pattern
    return re.match(pattern, email, re.VERBOSE) is not None

def resize_image_to_fit_cell(image, max_width, max_height):
    width, height = image.size
    aspect_ratio = width / height

    if width > max_width:
        width = max_width
        height = int(width / aspect_ratio)

    if height > max_height:
        height = max_height
        width = int(height * aspect_ratio)

    return image.resize((width, height))

# Sanitize the file name to avoid invalid characters
def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '', filename)

def is_signature_drawn(signature):
    # Check if signature is None or an empty numpy array
    if signature is None:
        return False
    # Ensure it is a numpy array and has content
    if isinstance(signature, np.ndarray) and signature.size > 0:
        # Additional check: if the array is not just empty white pixels
        # Assuming white background is [255, 255, 255] in RGB
        if np.all(signature == 255):
            return False
        return True
    return False

# Function to generate a unique ID using the UUID API
def generate_unique_id():
    try:
        response = requests.get("https://www.uuidtools.com/api/generate/v1")
        if response.status_code == 200:
            unique_id = response.json()[0]
            return unique_id
        else:
            st.warning("Error generating unique ID, fallback to internal ID.")
            return "fallback_id"  # Use a fallback ID in case of API failure
    except Exception as e:
        st.warning(f"Error generating unique ID: {e}, using fallback.")
        return "fallback_id"  # Fallback to a safe default ID

# Function to replace placeholders in the Word document
def replace_placeholder(paragraphs, placeholder, value=None, image_path=None):
    """ Replace a placeholder with a value or an image in the Word document without clearing the whole paragraph. """
    placeholder_with_brackets = f'[{placeholder}]'
    
    for paragraph in paragraphs:
        if placeholder_with_brackets in paragraph.text:
            # Split the paragraph at the placeholder, keep the text before and after
            parts = paragraph.text.split(placeholder_with_brackets)
            if len(parts) == 2:  # Ensure we are splitting at the placeholder
                paragraph.clear()  # Clear only the placeholder text, keep the rest intact
                # Rebuild the paragraph, add the image or value
                paragraph.add_run(parts[0])  # Text before the placeholder
                
                if image_path:
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(2))  # Insert the image
                elif value:
                    paragraph.add_run(value)  # Insert the value instead of the placeholder
                
                paragraph.add_run(parts[1])  # Text after the placeholder


# Function to populate the Word document with form data
def populate_document(data, template_path, resized_image_path, save_directory="/"):
    try:
        # Ensure that the save directory exists
        if not os.path.exists(save_directory):
            os.makedirs(save_directory, exist_ok=True)
        
        # Generate a unique ID for the filled document
        unique_id = generate_unique_id()

        # Load the Word document template
        doc = Document(template_path)
        paragraphs = doc.paragraphs

        # Replace placeholders with form data based on the updated template
        replace_placeholder(paragraphs, 'Add FULL name', data['learner_name'])
        replace_placeholder(paragraphs, 'Email', data['learner_email'])
        replace_placeholder(paragraphs, 'Phone', data['learner_phone'])
        replace_placeholder(paragraphs, 'Select Date', data['shared_date'])  # Use the same date everywhere
        
        # Insert the signature image instead of text
        replace_placeholder(paragraphs, 'Signature here', image_path=resized_image_path)
        
        replace_placeholder(paragraphs, 'Type your name', data['learner_name'])
        replace_placeholder(paragraphs, 'Select Date', data['shared_date'])  # Use the same date for learner's signature
        replace_placeholder(paragraphs, 'Enter Full Name', data['parent_signature'])
        replace_placeholder(paragraphs, 'Select Date', data['shared_date'])  # Use the same date for parent's signature

        # Auto-generated fields
        replace_placeholder(paragraphs, 'Auto-generated ID', unique_id)
        replace_placeholder(paragraphs, 'Auto-captured IP Address', socket.gethostbyname(socket.gethostname()))
        replace_placeholder(paragraphs, 'Auto-captured Timestamp', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

        # Save the filled document in the current directory with a unique name
        filled_doc_path = f"Filled_Consent_Form_{unique_id}.docx"
        doc.save(filled_doc_path)

        return filled_doc_path

    except Exception as e:
        st.error(f"Error processing the document: {e}")
        return None


# Function to send the document via Outlook using Streamlit secrets
def send_email(file_path):
    try:
        # sender_email = 'email'
        # password = 'pass'
        sender_email = st.secrets["sender_email"]  
        password = st.secrets["sender_password"]  
        receiver_email = sender_email  # Send to the same email address
        smtp_server = "smtp.office365.com"
        smtp_port = 587

        # Check if the file exists before attempting to send
        if not os.path.exists(file_path):
            st.warning("File not found. Skipping email sending.")
            return

        # Create a multipart message
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = receiver_email
        msg['Subject'] = "Digital Media Release Consent Form Submission"

        # Email body
        body = "Please find the attached filled consent form."
        msg.attach(MIMEText(body, 'plain'))

        # Attach the document
        with open(file_path, "rb") as attachment:
            part = MIMEApplication(attachment.read(), _subtype="docx")
            part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(file_path)}"')
            msg.attach(part)

        # Setup the server and send the email
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()  # Secure the connection
        server.login(sender_email, password)
        server.sendmail(sender_email, receiver_email, msg.as_string())
        server.quit()

        st.success(f"Consent form submitted and sent to {receiver_email}.")

    except smtplib.SMTPException as smtp_error:
        st.error(f"SMTP error occurred: {smtp_error}")

    except FileNotFoundError as file_error:
        st.error(f"Error with file handling: {file_error}")

    except Exception as e:
        st.error(f"An error occurred while sending the email: {e}")

# Streamlit form for user input
st.title("Digital Media Release Consent Form")

# Form Content
########################################################################
st.markdown("""
## Consent and Data Protection Agreement

**Dear Learner,**

We are requesting your permission to use the video and audio footage you have provided, including testimonials, interviews, declarations, course activities, and other visual or audio content (hereafter referred to as "Media Footage"), for the purposes outlined below. Your feedback is invaluable, and we would like to share it publicly for specific purposes.

By completing and submitting this digital form, you confirm that:

### 1. Use of Media Footage:
I understand that Prevista Ltd (the "Business") wishes to use the Media Footage for the following purposes (the "Specified Purposes"):

- In presentations, specifically at conferences, seminars, educational workshops, webinars, and internal meetings.
- In promotional materials, including brochures, flyers, social media posts, digital content, digital advertisements, posters, banners, and email campaigns.
- In advertising goods or services on digital platforms such as Google Ads and Facebook Ads, print ads, TV and radio ads, and online video and audio advertisements on platforms like YouTube, Instagram, X, and others.
- On the Business' website, prevista.co.uk, partner websites, internal and external blog articles, and course-specific pages.
- In marketing communications, including but not limited to educational, commercial, promotional, and informational uses, in any media or format, now known or invented in the future, worldwide, without time limits.

### 2. Rights and Permissions:
- I grant the Business exclusive permission to use the Media Footage for the Specified Purposes.
- I understand that my image and voice may be edited, copied, modified, exhibited, published, or distributed, and I waive the right to inspect or approve the finished product.

### 3. Compensation:
- I waive the right to any royalties or other compensation arising from or relating to the use of the Media Footage.

### 4. Data Storage and Transfer:
- I consent to the Business storing copies of the Media Footage and/or my contact details on its database for the Specified Purposes or in case it needs to contact me.
- I consent to the Business storing and transferring the Media Footage and my contact details to locations outside of the UK or European Economic Area (EEA), particularly to the Philippines and other regions that serve as bases for sub-contractors of the Business, for the Specified Purposes.

### 5. Data Protection and Privacy:
- In accordance with the UK General Data Protection Regulation (UK GDPR) and the Data Protection Act 2018, I consent to the collection, use, and processing of my personal data, including my name, likeness, voice, and any other identifiable information in the Media Footage, by Prevista Ltd solely for the purposes outlined in this form.
- I understand that my personal data will be processed fairly and lawfully and will not be used for any purpose other than those stated above without my additional consent.
- I understand that my data will be stored securely by Prevista Ltd and will only be retained for as long as necessary to fulfil the purposes outlined in this form or as required by law.

### 6. Right to Withdraw Consent:
- I understand that I may withdraw my consent at any time by contacting Prevista Ltd at enquiries@prevista.co.uk. However, I acknowledge that any use of my Media Footage prior to my withdrawal will not be affected.
- I am aware that if I withdraw my consent, my personal data will be deleted or anonymised where possible, in compliance with data protection laws.

### 7. Access to Information:
- I understand that I have the right to request access to the personal data held about me, to rectify any inaccuracies, or to request the erasure of my data where appropriate.
""")

# Authorization
########################################################################
'-----------------------------------------------------------------------'
st.markdown("""
**Authorisation**: 
I confirm that I am at least 18 years of age or have the consent of my parent/guardian to participate. I have read and fully understand this consent and data protection agreement.
""")

# Data Input
########################################################################

# Feedback form fields
learner_name = st.text_input("Type your name", key="learner_name")
learner_email = st.text_input("Type your email", key="learner_email")
learner_phone = st.text_input("Type your phone", key="learner_phone")
# shared_date = st.date_input(
#         label="Date",  # Label for the field
#         min_value=date(1900, 1, 1),  # Minimum selectable date
#         max_value=date.today(),  # Maximum selectable date
#         help="Choose a date",  # Tooltip text
#         format='DD/MM/YYYY',
#         key="shared_date"
#     )
canvas_result = st_canvas(
    fill_color="rgba(255, 255, 255, 1)",  
    stroke_width=5,
    stroke_color="rgb(0, 0, 0)",  # Black stroke color
    background_color="white",  # White background color
    width=400,
    height=150,
    drawing_mode="freedraw",
    key="learner_signature",
)
learner_signature = canvas_result.image_data
# learner_full_name = st.text_input("Learner's Full Name", key="learner_full_name")
parent_signature = st.text_input("Parent/Guardian's Signature (Typed Name, if applicable)", key="parent_signature")
# Set today's date automatically and display it
st.session_state.signature_date = date.today().strftime("%d-%m-%Y")
st.write(f"Date: **{st.session_state.signature_date}**")
st.markdown("""
**Note**: By clicking "Submit," you confirm that you have read, understood, and agree to the terms and conditions stated above.
""")

# Submit button
if st.button("Submit", key="submit_button", disabled=st.session_state.submission_status):
    if (is_valid_email(learner_email)):
        if learner_name and learner_email and learner_phone:
            if is_signature_drawn(learner_signature):
                with st.spinner('Processing...'):
                    try:
                        # Validate form inputs
                        # if not learner_name:
                    
                        # Collect form data
                        form_data = {
                            'learner_name': learner_name,
                            'learner_email': learner_email,
                            'learner_phone': learner_phone,
                            # 'shared_date': shared_date.strftime("%d-%m-%y"),  # Use shared date
                            'shared_date': st.session_state.signature_date,
                            'parent_signature': parent_signature,
                        }

                        # Path to the Word template document
                        template_path = "resource/ph_digital_media_consent.docx"  # Adjust the template file path

                        # Signature:
                        safe_learner_name = learner_name.strip().replace(" ", "_").lower()
                        signature_path = f'signature_{sanitize_filename(safe_learner_name)}.png'            
                        resized_image_path = f'resized_signature_image_{sanitize_filename(safe_learner_name)}.png'
                        signature_image = PILImage.fromarray(learner_signature.astype('uint8'), 'RGBA')
                        signature_image.save(signature_path)
                        # Open and resize the image
                        print(f"Opening image file: {signature_path}")
                        resized_image = PILImage.open(signature_path)
                        print(f"Original image size: {resized_image.size}")
                        resized_image = resize_image_to_fit_cell(resized_image, 200, 50)
                        resized_image.save(resized_image_path)  # Save resized image to a file
                        print(f"Resized image saved to: {resized_image_path}")


                        # Populate the document
                        filled_doc_path = populate_document(form_data, template_path, resized_image_path)

                        # Send the document via email
                        if filled_doc_path:
                            send_email(filled_doc_path)
                            st.session_state.submission_status = True
                            # st.experimental_rerun()
                            # download button
                            try:
                                # file download button
                                with open(filled_doc_path, 'rb') as f:
                                    file_contents = f.read()
                                    st.download_button(
                                        label="Download Your Response",
                                        data=file_contents,
                                        file_name=filled_doc_path,
                                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                                    )    
                            except FileNotFoundError as file_error:
                                st.error(f"Error with file handling: {file_error}")
                    except Exception as e:
                        # Display the error message on the screen
                        st.error(f"An unexpected error occurred")
                        st.error('Restarting in 10 SECONDS. . .')
                        st.error(f"Please take screenshot of the following error and share with Developer: \n{str(e)}")
                        time.sleep(12)
                        st.experimental_rerun()
            else:
                st.error("Please Draw the signature!")
        else:
            st.error("Please fill in all required fields!")
    else:
        st.warning("Please enter valid email address!")

if st.session_state.submission_status:
    st.success("Thank you for submitting your consent.")

# streamlit run app.py
# Dev : https://linkedin.com/in/osamatech786